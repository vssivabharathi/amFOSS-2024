import os
import csv
from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler, CallbackQueryHandler,
    ConversationHandler, ContextTypes, filters
)
import requests
from docx import Document

# Load environment variables from .env file
load_dotenv()

TELEGRAM_BOT_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
GOOGLE_BOOKS_API_KEY = os.getenv('GOOGLE_BOOKS_API_KEY')

app = ApplicationBuilder().token(TELEGRAM_BOT_TOKEN).build()

# Initialize reading list
reading_list = []

# State definitions for conversation handler
GENRE, BOOK_NAME, ACTION = range(3)

# Start command handler
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Hello! Welcome to PagePal! I can help you find books. Use /help to see all commands.')

# Help command handler
async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    help_text = (
        "/start - Start the bot\n"
        "/book - Search for books by genre\n"
        "/preview - Get a preview link for a specific book\n"
        "/list - Manage your reading list\n"
        "/reading_list - View your reading list\n"
        "/help - Show this help message"
    )
    await update.message.reply_text(help_text)

# Book command handler
async def book(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Please enter the genre of the book you want to find:')
    return GENRE  # Set state to GENRE

async def genre_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    genre = update.message.text
    books = fetch_books_by_genre(genre)
    filename = f"{genre}_books.csv"

    # Save to CSV file
    with open(filename, 'w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(["Title", "Author", "Description", "Year", "Language", "Preview"])
        for book in books:
            writer.writerow(book)

    # Send the CSV file
    await update.message.reply_document(document=open(filename, 'rb'))
    os.remove(filename)
    return ConversationHandler.END

# Fetch books by genre using Google Books API
def fetch_books_by_genre(genre):
    url = f"https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={GOOGLE_BOOKS_API_KEY}"
    response = requests.get(url).json()
    books = []

    for item in response.get("items", []):
        volume_info = item.get("volumeInfo", {})
        title = volume_info.get("title", "N/A")
        authors = ", ".join(volume_info.get("authors", []))
        description = volume_info.get("description", "N/A")
        published_date = volume_info.get("publishedDate", "N/A")
        language = volume_info.get("language", "N/A")
        preview_link = volume_info.get("previewLink", "N/A")
        books.append([title, authors, description, published_date, language, preview_link])

    return books

# Preview command handler
async def preview(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Please enter the book name for which you need a preview link:')
    return BOOK_NAME  # Set state to BOOK_NAME

async def book_name_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    book_name = update.message.text
    preview_link = fetch_book_preview(book_name)
    
    if preview_link:
        await update.message.reply_text(f"Preview link for '{book_name}': {preview_link}")
    else:
        await update.message.reply_text(f"Sorry, no preview available for '{book_name}'.")
    return ConversationHandler.END

# Fetch preview link for a specific book
def fetch_book_preview(book_name):
    url = f"https://www.googleapis.com/books/v1/volumes?q={book_name}&key={GOOGLE_BOOKS_API_KEY}"
    response = requests.get(url).json()
    items = response.get("items", [])

    if items:
        return items[0].get("volumeInfo", {}).get("previewLink", None)
    return None

# Reading list command handler
async def reading_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add')],
        [InlineKeyboardButton("Delete a book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text('Manage your reading list:', reply_markup=reply_markup)

# Callback query handler for reading list management
async def handle_buttons(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    if query.data == 'add':
        await query.edit_message_text(text="Enter the book name to add:")
        context.user_data['action'] = 'add'
        return BOOK_NAME

    elif query.data == 'delete':
        await query.edit_message_text(text="Enter the book name to delete:")
        context.user_data['action'] = 'delete'
        return BOOK_NAME

    elif query.data == 'view':
        doc = Document()
        doc.add_heading('Reading List', level=1)
        for book in reading_list:
            preview_link = fetch_book_preview(book)
            doc.add_paragraph(f"{book} - Preview: {preview_link if preview_link else 'No preview available'}")
        filename = 'reading_list.docx'
        doc.save(filename)
        await context.bot.send_document(chat_id=update.effective_chat.id, document=open(filename, 'rb'))
        os.remove(filename)
        return ConversationHandler.END

async def manage_reading_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    action = context.user_data.get('action')
    book_name = update.message.text

    if action == 'add':
        reading_list.append(book_name)
        await update.message.reply_text(f"Added '{book_name}' to the reading list.")
    elif action == 'delete':
        if book_name in reading_list:
            reading_list.remove(book_name)
            await update.message.reply_text(f"Deleted '{book_name}' from the reading list.")
        else:
            await update.message.reply_text(f"'{book_name}' is not in the reading list.")
    return ConversationHandler.END

# Main function to run the bot
def main():
    app = ApplicationBuilder().token(TELEGRAM_BOT_TOKEN).build()

    # Define the conversation handler
    conv_handler = ConversationHandler(
        entry_points=[
            CommandHandler('book', book),
            CommandHandler('preview', preview),
            CommandHandler('reading_list', reading_list)
        ],
        states={
            GENRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, genre_input)],
            BOOK_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, book_name_input)],
            ACTION: [MessageHandler(filters.TEXT & ~filters.COMMAND, manage_reading_list)],
        },
        fallbacks=[CommandHandler('help', help_command)]
    )

    # Add handlers to the application
    app.add_handler(CommandHandler('start', start))
    app.add_handler(CommandHandler('help', help_command))
    app.add_handler(conv_handler)
    app.add_handler(CallbackQueryHandler(handle_buttons))

    print("Bot is running...")
    app.run_polling()

if __name__ == '__main__':
    main()
